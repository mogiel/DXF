import Tag as Tag
from docx import Document
from docx.image import image


def run(doc: Document):
    monkey()
    para = doc.add_paragraph()
    pic = para.add_run().add_picture("Freesample.svg")
    compose_asvg(pic)


def compose_asvg(pic) -> None:
    embed = '{%s}embed' % (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships")
    url = "http://schemas.microsoft.com/office/drawing/2016/SVG/main"
    url14 = "http://schemas.microsoft.com/office/drawing/2010/main"
    etree.register_namespace("asvg", url)
    etree.register_namespace("a14", url14)

    ablip = pic._inline[-1][-1][-1][1][0]
    elst = OxmlElement('a:extLst')
    ablip.append(elst)
    embed_id = ablip.attrib.get(embed, "")
    del (ablip.attrib[embed])

    aext = OxmlElement('a:ext')
    elst.append(aext)
    aext.set("uri", "{28A0092B-C50C-407E-A947-70E740481C1C}")
    ldpi = etree.Element('{%s}useLocalDpi' % url14)
    aext.append(ldpi)
    ldpi.set("val", "0")

    aext = OxmlElement('a:ext')
    elst.append(aext)
    aext.set("uri", "{96DAC541-7B7A-43D3-8B79-37D633B846F1}")

    asvg = etree.Element('{%s}svgBlip' % url)
    aext.append(asvg)
    asvg.set(embed, embed_id)


def monkey() -> None:
    global default_factory
    if default_factory is not None:
        return
    default_factory = image._ImageHeaderFactory
    image._ImageHeaderFactory = _ImageHeaderFactory


def _ImageHeaderFactory(stream: IO[Text]) -> image.BaseImageHeader:
    stream.seek(0)
    buf = stream.read(32)
    if "<svg".encode("latin") in buf:
        stream.seek(0)
        return Svg.from_stream(stream)
    assert default_factory is not None
    return default_factory(stream)


class Svg(image.BaseImageHeader):
    @classmethod
    def from_stream(cls, stream: IO[Text]) -> '_SvgParser':
        """Return a |Svg| instance having header properties parsed from image
             in *stream*."""
        # this is dummy.
        return cls(100, 100, 96, 96)

    @property
    def content_type(self) -> Text:
        return "image/svg"
